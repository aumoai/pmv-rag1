import logging
import os

import PyPDF2
from docx import Document

logger = logging.getLogger(__name__)


class FileParser:
    """
    Parser for extracting text from various file formats
    """

    def __init__(self):
        self.supported_extensions: set[str] = {
            ".pdf",
            ".docx",
            ".txt",
            ".png",
            ".jpg",
            ".jpeg",
            ".webp",
        }

    async def extract_text(self, file_path: str) -> str | bytes:
        """
        Extract text from a file based on its extension
        """
        try:
            file_extension = os.path.splitext(file_path)[1].lower()

            if file_extension not in self.supported_extensions:
                raise ValueError(f"Unsupported file type: {file_extension}")

            if file_extension == ".pdf":
                return await self._extract_from_pdf(file_path)
            elif file_extension == ".docx":
                return await self._extract_from_docx(file_path)
            elif file_extension == ".txt":
                return await self._extract_from_txt(file_path)
            elif file_extension in {".png", ".jpg", ".jpeg", ".webp"}:
                return await self._extract_from_image(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")

        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise

    async def _extract_from_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF file
        """
        try:
            text_content: list[str] = []

            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)

                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text.strip():
                        text_content.append(text)

            extracted_text = "\n\n".join(text_content)
            logger.info(f"Extracted text from PDF: {len(extracted_text)} characters")
            return extracted_text

        except Exception as e:
            logger.error(f"Error extracting from PDF {file_path}: {str(e)}")
            raise

    async def _extract_from_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX file
        """
        try:
            doc = Document(file_path)
            text_content: list[str] = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))  # pyright: ignore[reportUnknownMemberType, reportUnknownArgumentType]

            extracted_text = "\n\n".join(text_content)
            logger.info(f"Extracted text from DOCX: {len(extracted_text)} characters")
            return extracted_text

        except Exception as e:
            logger.error(f"Error extracting from DOCX {file_path}: {str(e)}")
            raise

    async def _extract_from_txt(self, file_path: str) -> str:
        """
        Extract text from TXT file
        """
        try:
            with open(file_path, encoding="utf-8") as file:
                text_content = file.read()

            logger.info(f"Extracted text from TXT: {len(text_content)} characters")
            return text_content

        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, encoding="latin-1") as file:
                    text_content = file.read()
                logger.info(f"Extracted text from TXT (latin-1): {len(text_content)} characters")
                return text_content
            except Exception as e:
                logger.error(f"Error extracting from TXT {file_path}: {str(e)}")
                raise
        except Exception as e:
            logger.error(f"Error extracting from TXT {file_path}: {str(e)}")
            raise

    async def _extract_from_image(self, file_path: str) -> bytes:
        try:
            with open(file_path, "rb") as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error extracting from image {file_path}: {str(e)}")
            raise

    def get_supported_extensions(self) -> set[str]:
        """
        Get list of supported file extensions
        """
        return self.supported_extensions.copy()
